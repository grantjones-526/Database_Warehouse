"""Generate the project report as a .docx matching the formatting of the sample report."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

FONT = "Times New Roman"
BODY_SIZE = 12
HEADING_SIZE = 12


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(HEADING_SIZE)
    return p


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(BODY_SIZE)
    return p


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(BODY_SIZE)
    return p


doc = Document()

# Set default margins (1 inch all around)
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# --- Header ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Project Report")
run.bold = True
run.font.name = FONT
run.font.size = Pt(BODY_SIZE)

add_meta(doc, "Grant Jones")
add_meta(doc, "4/26/2026")
add_meta(doc, "CS456G")

# --- Project Description ---
add_heading(doc, "Project Description")

add_body(doc,
    "This project is a Financial Market Intelligence Data Warehouse — a multi-database "
    "system that collects, transforms, and analyzes stock market data through a star schema "
    "architecture and exposes an interactive Streamlit dashboard for both structured OLAP "
    "queries and natural language analysis. The system pulls daily stock price history from "
    "the yfinance API, financial news articles from the NewsAPI, and macroeconomic indicators "
    "from the FRED API, storing each dataset in the most appropriate database technology: "
    "MySQL for structured price data, MongoDB for unstructured news documents, and PostgreSQL "
    "for the analytical warehouse layer."
)

add_body(doc,
    "The dashboard provides two modes of interaction. The first is a set of pre-built OLAP "
    "queries covering sector performance, momentum analysis, period-over-period comparisons, "
    "news lead/lag impact, and a multi-factor stock screener. The second is a natural language "
    "interface powered by the Anthropic Claude API, which allows users to either ask plain-English "
    "analytical questions — translated directly into SQL — or request stock price forecasts, which "
    "trigger an ARIMA time-series model and return a plain-English decision recommendation "
    "enriched with recent news context from MongoDB."
)

# --- Functions Implemented ---
add_heading(doc, "Functions Implemented")

add_body(doc,
    "The project relies on three ETL scripts and two core application modules. The first ETL "
    "script, etl_stocks.py, downloads daily OHLCV (open, high, low, close, volume) price data "
    "for five tracked tickers — AAPL, MSFT, GOOGL, AMZN, and TSLA — using yfinance and loads "
    "it into a MySQL database with INSERT IGNORE to prevent duplicate rows across re-runs. "
    "The second script, etl_news.py, queries the NewsAPI across thirty targeted search terms, "
    "tags each article with matching tickers and financial topics using keyword matching, and "
    "stores the results in a MongoDB collection with a unique index on URL to deduplicate articles."
)

add_body(doc,
    "The third ETL script, etl_warehouse.py, builds the PostgreSQL star schema by generating "
    "the dim_date dimension across the full date range through today, copying asset and sector "
    "metadata into dim_asset and dim_sector, pre-loading MongoDB news counts into memory, and "
    "merging all data into the fact_market_data table using an upsert — ON CONFLICT DO UPDATE — "
    "so that news_count values are refreshed on each run without duplicating price records. "
    "After the fact table is populated, the script computes 30-day moving averages and "
    "volatility (standard deviation of daily returns) using PostgreSQL window functions and "
    "writes the results back into the ma_30 and vol_30 columns."
)

add_body(doc,
    "The application layer is split into two modules. The query_handler.py module in the claude "
    "package serves as the natural language routing layer: it inspects the user's question for "
    "buy or forecast intent using keyword matching, and routes to either the ARIMA prediction "
    "path or the SQL generation path. For SQL questions, it sends a compact schema description "
    "to Claude Haiku and extracts the resulting query. The regression.py module implements the "
    "ARIMA(5,1,0) time-series model using statsmodels: it fetches the last 90 trading days of "
    "close prices from PostgreSQL, fits the model, produces a 5-day forecast with an 80% "
    "confidence interval, and returns a compact summary that is combined with recent MongoDB "
    "headlines before being sent to Claude Haiku for narration."
)

# --- Technical Details ---
add_heading(doc, "Technical Details")

add_body(doc,
    "The system is built primarily with Python and uses three distinct database technologies "
    "to demonstrate a polyglot persistence architecture. MySQL 8 stores the operational stock "
    "price data in a normalized daily_prices table. MongoDB 7 stores financial news articles "
    "as documents with flexible schema, allowing ticker and topic arrays to vary per article. "
    "PostgreSQL 16 hosts the analytical warehouse in a star schema consisting of dim_date, "
    "dim_asset, dim_sector, and fact_market_data. All three databases are containerized and "
    "managed via Docker Compose, making the system fully portable and reproducible. The start.py "
    "entry point orchestrates container startup, dependency installation, ETL execution, and "
    "dashboard launch in a single command."
)

add_body(doc,
    "Key Python libraries include yfinance for market data retrieval, pymongo and "
    "mysql-connector-python for database access, psycopg2-binary for PostgreSQL, pandas for "
    "data manipulation, statsmodels for ARIMA modeling, and the Anthropic Python SDK for LLM "
    "integration. The dashboard is built with Streamlit, which provides interactive widgets, "
    "metric cards, data tables, and bar charts without requiring a separate frontend framework. "
    "Development was done in Visual Studio Code with the project running entirely on a local "
    "machine through Docker and Python's built-in virtual environment."
)

# --- Highlighted Features ---
add_heading(doc, "Highlighted Features")

add_body(doc,
    "A standout feature of this project is the decision layer built around the natural language "
    "interface. When a user asks a question such as 'Should I buy Apple stock next week?', the "
    "system does not simply retrieve data — it produces a structured analytical decision. The "
    "ARIMA(5,1,0) model is fitted on the last 90 trading days of closing prices pulled from "
    "PostgreSQL. The model differencing parameter d=1 removes the upward price trend to achieve "
    "stationarity, while the autoregressive parameter p=5 captures weekly autocorrelation patterns. "
    "The forecast produces tomorrow's predicted price, a 5-day predicted price, and an 80% "
    "confidence interval, all of which are surfaced as metric cards in the dashboard alongside "
    "the AIC model quality score."
)

add_body(doc,
    "The forecast is further enriched by pulling the five most recent news headlines for the "
    "queried ticker from MongoDB and passing them alongside the model output to Claude Haiku. "
    "The LLM receives only a compact summary of approximately 100 tokens — model statistics and "
    "headlines — and returns a 2-to-3 sentence plain-English recommendation that identifies the "
    "key trend direction, acknowledges the model's confidence interval and limitations, and "
    "highlights any relevant sentiment from recent news. This approach keeps API token usage "
    "minimal while producing contextually grounded output that draws on all three databases "
    "simultaneously, demonstrating the value of the multi-database warehouse design."
)

# --- Discussion ---
add_heading(doc, "Discussion")

add_body(doc,
    "Throughout the project, the three-database architecture proved to be both its greatest "
    "strength and its most significant engineering challenge. Storing data in the technology "
    "best suited to its structure — relational for prices, document for news, columnar star "
    "schema for analytics — made each individual query layer cleaner and more efficient. "
    "However, it required careful coordination during the warehouse ETL step to ensure that "
    "date formats aligned across MySQL, MongoDB, and PostgreSQL. A key issue encountered was "
    "that MongoDB news articles were dated in 2026 while the MySQL price history originally "
    "ended at 2025-12-31, causing news_count to remain zero across the entire fact table "
    "until both the stock ETL end date and the dim_date generation range were updated to "
    "use today's date dynamically."
)

add_body(doc,
    "The choice of ARIMA over simpler regression for price forecasting was justified by the "
    "nature of financial time series. Linear regression on time produces a trend line that "
    "ignores autocorrelation — the fact that yesterday's price is the strongest predictor of "
    "today's. ARIMA's differencing step makes the series stationary, and the autoregressive "
    "terms capture short-term momentum. The ARIMA(5,1,0) order follows a standard starting "
    "configuration for daily price series: one level of differencing and five autoregressive "
    "lags corresponding to one trading week. The AIC score is surfaced in the dashboard so "
    "users can assess model fit quality alongside the forecast."
)

add_body(doc,
    "There are several alternative approaches that could improve the system. A production "
    "deployment would benefit from a news source with historical coverage, such as a paid "
    "financial news API, to populate news_count across the full five-year date range rather "
    "than only the past month. The ARIMA model could be extended to ARIMAX or a gradient "
    "boosting model that incorporates volatility, news count, and moving average as features, "
    "taking fuller advantage of the warehouse's multi-dimensional data. Overall, the project "
    "demonstrates how a polyglot persistence architecture combined with an LLM interface can "
    "transform a standard data warehouse into an accessible analytical tool that bridges "
    "structured query results and natural language decision support."
)

doc.save("Project Report - Data Warehouse.docx")
print("Report saved to: Project Report - Data Warehouse.docx")
